from docx import Document
import docx
import requests
import json
from bs4 import BeautifulSoup
import requests
import pandas as pd
from pandas import Series, DataFrame
from ipywidgets import FloatProgress
from time import sleep
from IPython.display import display
import re
import pickle
import lxml
from textblob import TextBlob
from autocorrect import spell
import nltk
from nltk.probability import FreqDist
from nltk.corpus import webtext
import matplotlib.pyplot as plt
import speech_recognition as sr
import goslate


def data_from_word_files():
    document = open('demo.docx', 'rb')
    doc = Document(document)
    # documents = doc.Document(document)
    docu = ""
    for paragraph in doc.paragraphs:
        docu+=paragraph.text

def data_from_json():
    r = requests.get('https://quotes.rest/qod.json')
    res = r.json()
    print(json.dumps(res, indent = 4))

def data_from_scrap():
    url = 'https://www.imdb.com/chart/top?ref_=nv_mv_250_6'
    result = requests.get(url)
    content = result.content
    soup = BeautifulSoup(content,'lxml')
    summary = soup.find('div', {'class': 'article'})

    moviename = []
    cast = []
    description = []
    rating = []
    ratingoutof = []
    year = []
    genre = []
    movielength = []
    rot_audscore = []
    rot_avgrating = []
    rot_users = []

    rgx = re.compile('[%s]' % '()')
    f = FloatProgress(min=0, max=250)
    for row, i in zip(summary.find('table').findAll('tr'), range(len(summary.find('table').findAll('tr')))):
        for sitem in row.findAll('span', {'class': 'secondaryInfo'}):
            s = sitem.find(text=True)
            year.append(rgx.sub("", s))
            for ritem in row.findAll('td', {'class': 'ratingColumn imdbRating'}):
                for iget in ritem.findAll('strong'):
                    rating.append(iget.find(text=True))
                    ratingoutof.append(iget.get('title').split(' ', 4)[3])
        for item in row.findAll('td', {'class': 'titleColumn'}):
            for href in item.findAll('a', href=True):
                moviename.append(href.find(text=True))
    # List to pandas series
    moviename = Series(moviename)
    rating = Series(rating)
    ratingoutof = Series(ratingoutof)
    year = Series(year)

    imdb_df = pd.concat([moviename, year, rating, ratingoutof], axis=1)
    imdb_df.columns = ['moviename','year','imdb_rating','imdb_ratingbasedon']
    imdb_df['rank'] = imdb_df.index + 1
    print(imdb_df.head(1))
    imdb_df.to_csv("moviedata.csv")


def correct_spelling():
    text = ['Introduction to NLP', 'It is likely to be usefuel, to people ','Machine learning is the new electrcity', 'R is good langauage','I like this book','I want more books like this']
    # convert list to dataframe
    df = pd.DataFrame({'tweet': text})
    correct=df['tweet'].apply(lambda x: str(TextBlob(x).correct()))
    # print(correct)
    # b = TextBlob(df['tweet'])
    # print(b.correct())

    #auto correct
    # print(spell(u'sirvice'))

    #compute frequency of word
    # nltk.download('webtext')
    # wt_words = webtext.words('firefox.txt')
    lst=[]
    with open('paragraph.txt', 'r') as f:
        lst= f.read().split()

    frequency_dist = nltk.FreqDist(lst)
    sorted_frequency_dist = sorted(frequency_dist, key=frequency_dist.__getitem__, reverse = True)
    # print(sorted_frequency_dist)
    #words only if their frequency is greater than 3.
    large_words = dict([(k, v) for k, v in frequency_dist.items() if len(k) > 3])
    frequency_dist = nltk.FreqDist(large_words)
    frequency_dist.plot(50, cumulative=False)


def speech_to_text():
    r = sr.Recognizer()
    with sr.Microphone() as source:
        print("Say something!")
        audio = r.listen(source)
        print(r.recognize_google(audio))

    # r = sr.Recognizer()
    # with sr.Microphone() as source:
    #     r.adjust_for_ambient_noise(source)
    #     print("listening...")
    #     audio = r.listen(source)
    #
    #     try:
    #         str = r.recognize_google(audio)
    #         print(str)
    #     except:
    #         print("some error occurred!")

    # r = sr.Recognizer()
    # with sr.Microphone() as source:
    #     print("Please say something")
    #     audio = r.listen(source)
    #     r.adjust_for_ambient_noise(source, duration=1)
    #     print("Time over, thanks")
    # try:
    #     print("I think you said: " + r.recognize_google(audio))
    # except LookupError:
    #     print("Could not understand audio")

        # pass

def translate_sppech():
    text = "Bonjour le monde"
    gs = goslate.Goslate()
    translate = gs.translate(text,'ne')
    print(translate)


def creat_word_document():
    doc = docx.Document()

    doc.add_heading('Heading for the document', 0)

    doc_para = doc.add_paragraph('Your paragraph goes here, ')

    doc_para.add_run('hey there, bold here').bold = True
    doc_para.add_run(', and ')
    doc_para.add_run('these words are italic').italic = True

    # add a page break to start a new page
    doc.add_page_break()

    doc.add_heading('Heading level 2', 2)

    #add picture
    # doc.add_picture('path_to_picture')

    doc.save('demo.docx')


def main():
    # creat_word_document()
    # data_from_word_files()
    # data_from_json()
    # data_from_scrap()
    # correct_spelling()
    # speech_to_text()
    translate_sppech()


if __name__ == '__main__':
    main()
